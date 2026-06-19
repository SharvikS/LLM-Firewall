"""
File & image text extraction for the browser DLP extension's /scan-file path.

The browser extension can attach files and images to a chat (a leak path the
composer-text scan never sees). This module turns an uploaded file's raw bytes
into plain text so the SAME PII / secret / injection / code-leak detectors that
back /scan can run over it.

Dispatch is by extension + MIME type. Every backend is import-guarded and
best-effort: a missing optional dependency (PyMuPDF, python-docx, openpyxl, an
OCR engine) degrades that one file type to "unsupported" rather than breaking
the endpoint. Text-shaped files (code, csv, json, markdown, plain text) need no
dependency at all and always work.

Returns an ExtractResult so the caller can distinguish "file" from "image"
(for the dashboard breakdown) and tell genuine emptiness from an extractor that
wasn't available (so strict mode can fail closed on the latter).
"""

import concurrent.futures
import io
import logging
import os
import threading
from dataclasses import dataclass, field

logger = logging.getLogger("extract")

# Hard cap on extracted text fed to the detectors. Chat attachments can be large
# (a multi-MB PDF), but the detectors only need a representative slice — and
# Presidio is O(n) per recognizer, so an unbounded paste would stall the worker.
MAX_EXTRACT_CHARS = int(os.getenv("DLP_MAX_EXTRACT_CHARS", "200000"))

# Wall-clock budget for a single image's OCR. Both servers that call this run
# multi-threaded (HTTP ThreadingHTTPServer, gRPC ThreadPoolExecutor), so a
# pathological image must not tie up a worker indefinitely. On timeout the image
# comes back unsupported (and, being an image, fails OPEN — never a false block).
OCR_TIMEOUT_S = float(os.getenv("DLP_OCR_TIMEOUT_S", "20"))

# Extensions we treat as already-plain-text (decode directly, no library).
_TEXT_EXTS = {
    ".txt", ".text", ".log", ".md", ".markdown", ".rst",
    ".csv", ".tsv", ".json", ".jsonl", ".ndjson", ".yaml", ".yml", ".toml",
    ".ini", ".cfg", ".conf", ".env", ".properties",
    ".xml", ".html", ".htm", ".svg",
    ".py", ".js", ".jsx", ".ts", ".tsx", ".go", ".rs", ".java", ".kt",
    ".c", ".h", ".cc", ".cpp", ".hpp", ".cs", ".rb", ".php", ".swift",
    ".sh", ".bash", ".zsh", ".ps1", ".sql", ".r", ".scala", ".pl", ".lua",
    ".tf", ".dockerfile", ".gradle", ".makefile",
}

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif"}


@dataclass
class ExtractResult:
    """Outcome of an extraction attempt.

    kind        : "file" | "image" (drives the dashboard breakdown)
    text        : extracted text (possibly empty)
    supported   : did we have a backend for this type at all?
    extractor   : which path produced the text (for diagnostics)
    truncated   : was the text capped at MAX_EXTRACT_CHARS?
    error       : human-readable reason when extraction could not run
    """
    kind: str = "file"
    text: str = ""
    supported: bool = True
    extractor: str = "none"
    truncated: bool = False
    error: str = ""
    meta: dict = field(default_factory=dict)


def _ext(filename: str) -> str:
    name = (filename or "").lower().strip()
    if "." not in name:
        # Files like "Dockerfile" / "Makefile" carry the type in the stem.
        stem = name.rsplit("/", 1)[-1]
        return "." + stem if stem in ("dockerfile", "makefile") else ""
    return name[name.rfind("."):]


def _cap(text: str) -> tuple[str, bool]:
    if len(text) > MAX_EXTRACT_CHARS:
        return text[:MAX_EXTRACT_CHARS], True
    return text, False


def _is_image(ext: str, content_type: str) -> bool:
    return ext in _IMAGE_EXTS or (content_type or "").startswith("image/")


# Magic-byte signatures for the image formats we OCR. The caller-supplied
# filename / content_type are untrusted (a direct API call or a compromised
# extension can spoof them), and "image" is a privileged classification — images
# fail OPEN even under strict policy. So we only grant image treatment when the
# bytes themselves prove it, which also means a real image mislabelled as
# ".bin"/octet-stream is still recognised and OCR'd.
_IMAGE_MAGIC = (
    b"\x89PNG\r\n\x1a\n",   # PNG
    b"\xff\xd8\xff",        # JPEG
    b"GIF87a", b"GIF89a",   # GIF
    b"BM",                  # BMP
    b"II*\x00", b"MM\x00*", # TIFF (little- / big-endian)
)


def _sniff_image(data: bytes) -> bool:
    """True iff the leading bytes match a known image signature."""
    head = data[:16]
    if any(head.startswith(sig) for sig in _IMAGE_MAGIC):
        return True
    # WEBP is a RIFF container: "RIFF"....(size)...."WEBP".
    return head[:4] == b"RIFF" and head[8:12] == b"WEBP"


# ── Per-format extractors (each import-guarded, returns text or raises) ───────

def _extract_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF  # noqa: PLC0415
    parts = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text())
            if sum(len(p) for p in parts) > MAX_EXTRACT_CHARS:
                break
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx  # noqa: PLC0415
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    import openpyxl  # noqa: PLC0415
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
            if sum(len(p) for p in parts) > MAX_EXTRACT_CHARS:
                break
    wb.close()
    return "\n".join(parts)


def _extract_text_bytes(data: bytes) -> str:
    # Try strict UTF-8 first, then a lossy decode so a stray byte never loses the
    # whole file (we'd rather scan slightly-garbled text than miss a secret).
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="replace")


# OCR backends are lazy + cached: importing torch/easyocr is expensive, so we
# only pay it on the first image and keep the reader resident afterwards. Both
# callers are multi-threaded, so a lock guards the one-time reader construction
# (EasyOCR's Reader() is not safe to enter concurrently — two threads racing it
# would each download/load the weights and could corrupt the shared cache).
_ocr_reader = None
_ocr_backend = None  # "easyocr" | "tesseract" | "unavailable"
_ocr_lock = threading.Lock()


def _easyocr_reader():
    """Build the EasyOCR reader once, under a lock, and cache it.

    Double-checked locking: the common path (reader already built) takes no lock;
    only the first caller constructs it while the rest wait, so concurrent image
    scans can never build two readers or race the weight download.
    """
    global _ocr_reader
    if _ocr_reader is not None:
        return _ocr_reader
    with _ocr_lock:
        if _ocr_reader is None:
            import contextlib  # noqa: PLC0415
            import easyocr  # noqa: PLC0415
            # English only by default; GPU off for a CPU-only deployment. The
            # first construction downloads detector/recogniser weights and prints
            # a progress bar with block glyphs that crash a non-UTF-8 console
            # (Windows cp1252), so we silence its stdout/stderr during build and
            # pass verbose=False to mute the per-call progress bar. Set
            # DLP_OCR_LANGS (comma-separated, e.g. "en,fr,de") for multilingual
            # documents — EasyOCR supports 80+ languages.
            langs = [l.strip() for l in os.getenv("DLP_OCR_LANGS", "en").split(",") if l.strip()] or ["en"]
            with open(os.devnull, "w") as devnull, \
                    contextlib.redirect_stdout(devnull), contextlib.redirect_stderr(devnull):
                _ocr_reader = easyocr.Reader(langs, gpu=False, verbose=False)
    return _ocr_reader


def _ocr_image_inner(data: bytes) -> str:
    global _ocr_backend

    # Prefer tesseract when a binary is actually installed (fast, light); else
    # fall back to easyocr (torch-based, no system binary, downloads models once).
    if _ocr_backend in (None, "tesseract"):
        try:
            import pytesseract  # noqa: PLC0415
            from PIL import Image  # noqa: PLC0415
            # Probe for the binary; pytesseract raises if it's missing.
            pytesseract.get_tesseract_version()
            _ocr_backend = "tesseract"
            img = Image.open(io.BytesIO(data))
            return pytesseract.image_to_string(img)
        except Exception as exc:  # noqa: BLE001
            logger.debug("tesseract OCR unavailable (%s) — trying easyocr", exc)

    try:
        reader = _easyocr_reader()
        _ocr_backend = "easyocr"
        lines = reader.readtext(data, detail=0, paragraph=True)
        return "\n".join(lines)
    except Exception as exc:  # noqa: BLE001
        _ocr_backend = "unavailable"
        raise RuntimeError(f"no OCR backend available: {exc}") from exc


def _ocr_image(data: bytes) -> str:
    """Run OCR with a wall-clock budget so a pathological image can't pin a
    worker thread forever. The OCR work itself is CPU-bound and not cancellable,
    so on timeout we surrender the request (raise) and let the daemon thread
    finish on its own rather than blocking the caller."""
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
        future = pool.submit(_ocr_image_inner, data)
        try:
            return future.result(timeout=OCR_TIMEOUT_S)
        except concurrent.futures.TimeoutError as exc:
            raise RuntimeError(f"OCR timed out after {OCR_TIMEOUT_S:g}s") from exc


def extract_text(filename: str, content_type: str, data: bytes) -> ExtractResult:
    """Extract scannable text from an uploaded file's raw bytes.

    Never raises for an *expected* failure (missing optional dep, unknown type,
    OCR backend absent): those come back as supported=False with an `error`, so
    the caller can decide policy (e.g. strict mode blocks; lenient allows). Only
    truly unexpected bugs propagate.
    """
    ext = _ext(filename)
    ct = (content_type or "").lower()

    # ── Images → OCR ──────────────────────────────────────────────────────────
    # Trust the bytes, not the label: route to OCR only when a magic-byte sniff
    # confirms a real image. This both rescues mislabelled images and denies the
    # privileged "image" (fail-open) treatment to a non-image that merely claims
    # an image content_type/extension — such a file falls through to the document
    # / binary path below, where strict mode can fail closed on it.
    if _sniff_image(data) or (_is_image(ext, ct) and not data):
        try:
            text = _ocr_image(data)
            text, truncated = _cap(text)
            return ExtractResult(kind="image", text=text, supported=True,
                                 extractor=f"ocr:{_ocr_backend}", truncated=truncated,
                                 meta={"ocr_attempted": True, "ocr_failed": False,
                                       "ocr_backend": _ocr_backend})
        except Exception as exc:  # noqa: BLE001
            return ExtractResult(kind="image", supported=False, extractor="ocr",
                                 error=str(exc),
                                 meta={"ocr_attempted": True, "ocr_failed": True})

    # ── Documents → format-specific extractor ────────────────────────────────
    handlers = []
    if ext == ".pdf" or "pdf" in ct:
        handlers = [("pdf", _extract_pdf)]
    elif ext == ".docx" or "wordprocessingml" in ct:
        handlers = [("docx", _extract_docx)]
    elif ext in (".xlsx", ".xlsm") or "spreadsheetml" in ct:
        handlers = [("xlsx", _extract_xlsx)]

    for name, fn in handlers:
        try:
            text, truncated = _cap(fn(data))
            return ExtractResult(kind="file", text=text, supported=True,
                                 extractor=name, truncated=truncated)
        except ImportError as exc:
            return ExtractResult(kind="file", supported=False, extractor=name,
                                 error=f"{name} support not installed ({exc})")
        except Exception as exc:  # noqa: BLE001
            return ExtractResult(kind="file", supported=False, extractor=name,
                                 error=f"{name} extraction failed: {exc}")

    # ── Plain-text-shaped files (and text/* MIME) → decode directly ──────────
    if ext in _TEXT_EXTS or ct.startswith("text/") or "json" in ct or "xml" in ct:
        text, truncated = _cap(_extract_text_bytes(data))
        return ExtractResult(kind="file", text=text, supported=True,
                             extractor="text", truncated=truncated)

    # ── Unknown binary: try a best-effort UTF-8 sniff, else give up cleanly ──
    try:
        sniff = data.decode("utf-8")
        text, truncated = _cap(sniff)
        return ExtractResult(kind="file", text=text, supported=True,
                             extractor="text-sniff", truncated=truncated)
    except UnicodeDecodeError:
        return ExtractResult(kind="file", supported=False, extractor="none",
                             error=f"unsupported file type ({ext or ct or 'unknown'})")
